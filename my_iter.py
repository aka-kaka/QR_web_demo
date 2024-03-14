"""
   No
"""
from glob import glob
from os import curdir, sep as os_sep
from shutil import move

import pandas as pd
from docx import Document


class IterFiles:
    """
    No Doc
    """
    def __init__(
            self,
            path_input=curdir + os_sep + "input" + os_sep,
            search_doc="doc",
    ):
        self._glob = [i for i in glob(f"{path_input}*.{search_doc}*")
                      if i != path_input+"QR_cashflows.xlsx"]

    def __iter__(self):
        self._glob = iter(self._glob)
        return self

    def __next__(self):

        return next(self._glob)

    def lendth(self):
        """
        No DOc
        """
        return len(list(self._glob))


class GetData():
    """
    No DOc
    """
    def __init__(
            self,
            file_extension: str,
            file_name=None,
    ):
        self.file_extension = file_extension
        self.doc_table = None
        self.excel_table = None
        self.file_name = file_name

    @staticmethod
    def move_file(file):
        """
        No DOC
        """

        tmp_file = file.split(os_sep)
        if len(tmp_file) > 1:
            tmp_file[-2] = "done"
            move(file, os_sep.join(tmp_file))
        else:
            move(file, f"{curdir}{os_sep}done{os_sep}{file}")

    def get_data(
            self):
        """
        No DOc
        """
        if not self.file_name:
            tmp = IterFiles(search_doc=self.file_extension)
        else:
            tmp = [self.file_name,]

        if "doc" in self.file_extension:
            if tmp.lendth() > 0:
                self.doc_table = {
                    file.split(os_sep)[-1].split('.')[0]:
                        self.get_all_tables_file(file)
                    for file in tmp
                }

                # print(self.doc_table)
                # print(self.doc_table["000000"])
                # print(len(self.doc_table))

        if "xls" in self.file_extension:
            if ((isinstance(tmp, list) and
                 len(tmp) > 0) or tmp.lendth()) > 0:
                self.excel_table = {
                    file.split(os_sep)[-1].split('.')[0]:
                        self.read_excel(file)
                    for file in tmp
                }

    def read_excel(self, file):
        """
        No DOc
        """
        df = pd.read_excel(file)
        self.move_file(file)
        return df.to_dict()

    def get_all_tables_file(self, file: str) -> dict:
        """
        No DOc
        """
        res = {}
        for table in Document(file).tables:
            tmp = self.get_data_from_tables(table=table)
            if tmp:
                res.update(tmp)
        self.move_file(file)
        return res if res else None

    def get_all_tables_table(self, tables):

        """
        NO Doc
        """

        res = {}
        for table in tables:
            tmp = self.get_data_from_tables(table)
            if tmp:
                res.update(tmp)
        return res

    def get_data_from_tables(self, table):

        """
        NO Doc
        """

        res = {}

        for row in table.rows:
            for cell in row.cells:
                if cell.tables:
                    tab = cell.tables
                    res.update(self.get_all_tables_table(tab))
            if len(row.cells) == 2:
                res.update(
                    {row.cells[0].text.strip(): row.cells[1].text.strip()})
        return res if res else None
